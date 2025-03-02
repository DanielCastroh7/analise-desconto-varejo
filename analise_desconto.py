import streamlit as st
import pandas as pd
import plotly.express as px
import requests
from docx import Document

# Configuração da página
st.set_page_config(page_title="Análise de Desconto", page_icon="💰", layout="wide")

# Título da aplicação
st.title("💰 Análise de Desconto para Varejistas")

# Subtítulo
st.markdown("""
**Descubra se é vantajoso oferecer um desconto no seu produto, compare estratégias, simule cenários e analise tendências de mercado.**
""")

# Inputs do usuário
st.sidebar.header("📊 Dados do Produto")
valor_produto = st.sidebar.number_input("Valor do Produto (R$)", min_value=0.0, value=37.0, step=1.0)
quantidade_vendida_hoje = st.sidebar.number_input("Quantidade Vendida Hoje", min_value=0, value=1000, step=1)
margem_lucro = st.sidebar.number_input("Margem de Lucro por Unidade (%)", min_value=0.0, max_value=100.0, value=20.0, step=1.0) / 100

st.sidebar.header("📉 Desconto e Custos de Marketing")
desconto = st.sidebar.number_input("Desconto Proposto (%)", min_value=0.0, max_value=100.0, value=5.0, step=1.0) / 100
cac_percentual = st.sidebar.number_input("Custo de Aquisição de Clientes (CAC) em % da Margem", min_value=0.0, max_value=100.0, value=1.0, step=1.0) / 100

st.sidebar.header("📈 Quantidade Estimada de Vendas")
quantidade_estimada = st.sidebar.number_input("Quantidade Estimada de Vendas com Desconto", min_value=0, value=1350, step=1)

# Função para calcular a receita e o lucro
def calcular_receita(valor_produto, quantidade, margem, desconto, cac_percentual):
    margem_atualizada = margem - desconto - (cac_percentual * margem)
    lucro_por_unidade = valor_produto * margem_atualizada
    receita_total = quantidade * lucro_por_unidade
    return receita_total, lucro_por_unidade, margem_atualizada

# Cálculo do cenário atual
receita_hoje, lucro_hoje, _ = calcular_receita(valor_produto, quantidade_vendida_hoje, margem_lucro, 0, 0)

# Cálculo do cenário com desconto
receita_desconto, lucro_desconto, margem_atualizada = calcular_receita(valor_produto, quantidade_estimada, margem_lucro, desconto, cac_percentual)

# Função para gerar o documento Word
def gerar_documento(valor_produto, quantidade_vendida_hoje, margem_lucro, desconto, cac_percentual, quantidade_estimada, receita_hoje, receita_desconto):
    doc = Document()
    doc.add_heading('Relatório de Análise de Desconto', 0)

    # Adicionar dados do produto
    doc.add_heading('Dados do Produto', level=1)
    doc.add_paragraph(f"Valor do Produto: R$ {valor_produto:.2f}")
    doc.add_paragraph(f"Quantidade Vendida Hoje: {quantidade_vendida_hoje} unidades")
    doc.add_paragraph(f"Margem de Lucro: {margem_lucro * 100:.2f}%")

    # Adicionar dados de desconto e CAC
    doc.add_heading('Desconto e Custos de Marketing', level=1)
    doc.add_paragraph(f"Desconto Proposto: {desconto * 100:.2f}%")
    doc.add_paragraph(f"Custo de Aquisição de Clientes (CAC): {cac_percentual * 100:.2f}%")

    # Adicionar quantidade estimada de vendas
    doc.add_heading('Quantidade Estimada de Vendas', level=1)
    doc.add_paragraph(f"Quantidade Estimada de Vendas com Desconto: {quantidade_estimada} unidades")

    # Adicionar análise básica
    doc.add_heading('Análise Básica', level=1)
    doc.add_paragraph(f"Receita Total Hoje: R$ {receita_hoje:,.2f}")
    doc.add_paragraph(f"Receita Total com Desconto: R$ {receita_desconto:,.2f}")

    # Adicionar conclusão
    doc.add_heading('Conclusão', level=1)
    if receita_desconto > receita_hoje:
        doc.add_paragraph("✅ **É vantajoso oferecer o desconto!**")
        doc.add_paragraph(f"O lucro adicional será de **R$ {receita_desconto - receita_hoje:,.2f}**.")
    else:
        doc.add_paragraph("❌ **Não é vantajoso oferecer o desconto.**")
        doc.add_paragraph(f"Você teria uma perda de **R$ {receita_hoje - receita_desconto:,.2f}**.")

    # Salvar o documento
    doc.save("relatorio_analise_desconto.docx")

# Botão para gerar o documento
if st.sidebar.button("📄 Gerar Documento"):
    gerar_documento(valor_produto, quantidade_vendida_hoje, margem_lucro, desconto, cac_percentual, quantidade_estimada, receita_hoje, receita_desconto)
    st.sidebar.success("✅ Documento gerado com sucesso!")
    with open("relatorio_analise_desconto.docx", "rb") as file:
        st.sidebar.download_button(label="⬇️ Baixar Documento", data=file, file_name="relatorio_analise_desconto.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Abas para análise
aba1, aba2, aba3 = st.tabs(["Análise Básica", "Comparação de Estratégias", "Análise de Mercado e Tendências"])

# Aba 1: Análise Básica
with aba1:
    st.header("📊 Análise Básica")
    st.markdown("""
    **Aqui você pode ver o impacto do desconto na receita e no lucro, considerando a quantidade estimada de vendas.**
    """)

    st.subheader("📌 Cenário Atual")
    st.write(f"- **Quantidade Vendida Hoje:** {quantidade_vendida_hoje} unidades")
    st.write(f"- **Lucro por Unidade:** R$ {lucro_hoje:.2f}")
    st.write(f"- **Receita Total Hoje:** R$ {receita_hoje:,.2f}")

    st.subheader("📌 Cenário com Desconto")
    st.write(f"- **Quantidade Estimada de Vendas:** {quantidade_estimada} unidades")
    st.write(f"- **Lucro por Unidade com Desconto:** R$ {lucro_desconto:.2f}")
    st.write(f"- **Receita Total com Desconto:** R$ {receita_desconto:,.2f}")

    # Conclusão
    st.header("🎯 Conclusão")
    if receita_desconto > receita_hoje:
        st.success(f"✅ **É vantajoso oferecer o desconto!**")
        st.write(f"O lucro adicional será de **R$ {receita_desconto - receita_hoje:,.2f}**.")
    else:
        st.error(f"❌ **Não é vantajoso oferecer o desconto.**")
        st.write(f"Você teria uma perda de **R$ {receita_hoje - receita_desconto:,.2f}**.")

# Aba 2: Comparação de Estratégias
with aba2:
    st.header("🏆 Comparação de Estratégias")
    st.markdown("""
    **Compare diferentes estratégias de desconto e veja qual é a mais vantajosa.**
    """)

    # Estratégias para comparação
    estrategias = [
        {"nome": "Sem Desconto", "desconto": 0.0, "cac_percentual": 0.0},
        {"nome": "Desconto Pequeno (5%)", "desconto": 0.05, "cac_percentual": 0.01},
        {"nome": "Desconto Médio (10%)", "desconto": 0.10, "cac_percentual": 0.02},
        {"nome": "Desconto Grande (15%)", "desconto": 0.15, "cac_percentual": 0.03},
        {"nome": "Desconto Agressivo (20%)", "desconto": 0.20, "cac_percentual": 0.04},
    ]

    # Cálculo dos cenários para cada estratégia
    resultados = []
    for estrategia in estrategias:
        receita, lucro_por_unidade, margem_atualizada = calcular_receita(
            valor_produto, quantidade_estimada, margem_lucro, estrategia["desconto"], estrategia["cac_percentual"]
        )
        vantajoso = receita > receita_hoje
        resultados.append({
            "Estratégia": estrategia["nome"],
            "Desconto": f"{estrategia['desconto'] * 100:.2f}%",
            "CAC": f"{estrategia['cac_percentual'] * 100:.2f}%",
            "Quantidade Estimada": quantidade_estimada,
            "Margem de Lucro": f"{margem_atualizada * 100:.2f}%",
            "Receita Total": f"R$ {receita:,.2f}",
            "Vantajoso": "✅ Sim" if vantajoso else "❌ Não"
        })

    # Tabela de comparação
    st.dataframe(pd.DataFrame(resultados))

    # Sugestão de Melhor Estratégia (excluindo "Sem Desconto")
    st.subheader("🎯 Melhor Estratégia com Desconto")
    estrategias_com_desconto = [r for r in resultados if r["Estratégia"] != "Sem Desconto"]
    if estrategias_com_desconto:
        melhor_estrategia = max(estrategias_com_desconto, key=lambda x: float(x["Receita Total"].replace("R$", "").replace(",", "").strip()))
        st.success(f"✅ **A melhor estratégia com desconto é: {melhor_estrategia['Estratégia']}**")
        st.write(f"- **Receita Total:** {melhor_estrategia['Receita Total']}")
        st.write(f"- **Margem de Lucro:** {melhor_estrategia['Margem de Lucro']}")
    else:
        st.warning("⚠️ **Nenhuma estratégia com desconto foi avaliada.**")

    # Gráfico de Melhores Estratégias
    st.subheader("📊 Gráfico de Melhores Estratégias")
    df_grafico = pd.DataFrame(resultados)

    # Limpeza e conversão da coluna "Receita Total"
    df_grafico["Receita Total"] = (
        df_grafico["Receita Total"]
        .str.replace("R$", "")  # Remove o símbolo "R$"
        .str.replace(".", "")   # Remove os pontos de milhar
        .str.replace(",", ".")  # Substitui a vírgula decimal por ponto
        .astype(float)          # Converte para float
    )

    # Criação do gráfico
    fig = px.bar(df_grafico, x="Estratégia", y="Receita Total", color="Vantajoso", title="Receita Total por Estratégia")
    st.plotly_chart(fig)
# Aba 3: Análise de Mercado e Tendências
# Aba 3: Análise de Mercado e Tendências
with aba3:
    st.header("🌍 Análise de Mercado e Tendências")
    st.markdown("""
    **Aqui você pode ver dados externos e tendências de mercado que podem impactar suas estratégias de desconto.**
    """)

    # Integração com API pública (exemplo: Alpha Vantage para dados financeiros)
    st.subheader("📈 Dados Externos (Exemplo: Taxa de Câmbio)")
    api_key = "VR8GHJQ3J1CFHT8T"  # Substitua pela sua chave da Alpha Vantage
    url = f"https://www.alphavantage.co/query?function=FX_DAILY&from_symbol=USD&to_symbol=BRL&apikey={api_key}"
    response = requests.get(url)
    data = response.json()

    if "Time Series FX (Daily)" in data:
        # Transforma os dados em um DataFrame
        df_cambio = pd.DataFrame(data["Time Series FX (Daily)"]).T
        df_cambio.index = pd.to_datetime(df_cambio.index)
        
        # Converte os valores para float
        df_cambio["4. close"] = df_cambio["4. close"].astype(float)
        #print(df_cambio["4. close"])
        
        # Exibe os dados
        st.write("**Taxa de Câmbio USD/BRL (Últimos 30 dias):**")
        st.line_chart(df_cambio["4. close"].tail(30))

        # Análise relacionada ao câmbio
        st.subheader("📌 Impacto do Câmbio na Estratégia")
        st.write("""
        - **Câmbio Alto (USD caro):** Pode aumentar os custos de produtos importados, reduzindo a margem de lucro.
        - **Câmbio Baixo (USD barato):** Pode reduzir os custos de produtos importados, aumentando a margem de lucro.
        """)
    else:
        st.warning("Não foi possível carregar os dados de câmbio. Verifique sua chave da API ou tente novamente mais tarde.")

    # Outros dados externos (exemplo: IBGE para dados econômicos)
    st.subheader("📊 Dados Econômicos (Exemplo: Inflação)")
    st.write("""
    - **Inflação Alta:** Pode reduzir o poder de compra dos consumidores, afetando as vendas.
    - **Inflação Baixa:** Pode aumentar o poder de compra dos consumidores, impulsionando as vendas.
    """)